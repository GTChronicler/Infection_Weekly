from __future__ import annotations

import csv
import math
import re
import sys
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from python_calamine import load_workbook


CONFIG_TEMPLATE = """# Infection Weekly report configuration
# 周报参数配置文件。请使用记事本、VS Code、Positron 等纯文本编辑器修改，不建议用 Excel 打开或保存。
# 空行和以 # 开头的说明行会被脚本忽略。
#
# [settings]
# date_Mon 使用 YYYYMMDD 格式，填写本次周报所描述周的周一日期。
#
# [overview_diseases]
# 填写“一、疫情概况”章节中需要点名比较的疾病。每行一个疾病名称，名称需与 config/dictionary.csv 中的 diseases 一致。
#
# [focus_diseases]
# 填写“二、重点疫情”章节中需要展开分析的疾病。每行一个疾病名称，名称需同时能在 A.xlsx 和 B.xlsx 中匹配。

[settings]
date_Mon={date_mon}

[overview_diseases]
{overview_diseases}

[focus_diseases]
{focus_diseases}
"""

SECTION_LABELS = [
    "（一）",
    "（二）",
    "（三）",
    "（四）",
    "（五）",
    "（六）",
    "（七）",
    "（八）",
    "（九）",
    "（十）",
    "（十一）",
]


class ConfigError(Exception):
    pass


@dataclass
class ReportConfig:
    date_mon: str
    overview_diseases: list[str]
    focus_diseases: list[str]


@dataclass
class ReportResult:
    docx_path: Path
    csv_path: Path
    paragraphs: list[str]
    table_rows: list[list[Any]]


def project_root() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parents[1]


def strip_spaces(value: Any) -> str:
    return "" if value is None else str(value).replace(" ", "")


def as_number(value: Any) -> float:
    if value is None or value == "":
        return 0.0
    if isinstance(value, (int, float)):
        if isinstance(value, float) and math.isnan(value):
            return 0.0
        return float(value)
    text = str(value).strip()
    if text in {"", "-", "NA", "nan"}:
        return 0.0
    return float(text)


def fmt_int(value: Any) -> str:
    return str(int(as_number(value)))


def fmt_pct(value: float, digits: int = 2) -> str:
    return f"{value:.{digits}f}"


def join_zh(values: Iterable[str], sep: str = "、") -> str:
    return sep.join(values)


def read_report_config(path: Path) -> ReportConfig:
    try:
        raw_lines = path.read_text(encoding="utf-8-sig").splitlines()
    except UnicodeDecodeError as exc:
        raise ConfigError(
            "配置文件不是 UTF-8 编码 / Config file is not UTF-8 encoded. "
            "请将 config/report_config.txt 保存为 UTF-8 / Please save config/report_config.txt as UTF-8."
        ) from exc

    lines = []
    for raw_line in raw_lines:
        line = raw_line.strip()
        if line and not line.startswith("#"):
            lines.append(line)

    settings: dict[str, str] = {}
    overview: list[str] = []
    focus: list[str] = []
    section: str | None = None

    for line in lines:
        if re.fullmatch(r"\[[^\]]+\]", line):
            section = line.strip("[]")
            continue

        if section is None:
            raise ConfigError(
                f"配置行未放在任何 section 下 / Config line is outside any section: {line}"
            )

        if section == "settings":
            key, sep, value = line.partition("=")
            if not sep or value.strip() == "":
                raise ConfigError(f"配置项格式无效 / Invalid settings line: {line}")
            settings[key.strip()] = value.strip()
        elif section == "overview_diseases":
            overview.append(line)
        elif section == "focus_diseases":
            focus.append(line)
        else:
            raise ConfigError(f"未知配置 section / Unknown config section: {section}")

    date_mon = settings.get("date_Mon")
    if date_mon is None:
        raise ConfigError("缺少 [settings] 中的 date_Mon / Missing date_Mon in [settings].")
    if not re.fullmatch(r"\d{8}", date_mon):
        raise ConfigError(
            "date_Mon 格式无效，请使用 YYYYMMDD，例如 20260518 / "
            "Invalid date_Mon in [settings]. Use YYYYMMDD, for example 20260518."
        )
    try:
        parsed = datetime.strptime(date_mon, "%Y%m%d")
    except ValueError as exc:
        raise ConfigError(
            "date_Mon 不是一个真实日期 / "
            "Invalid date_Mon in [settings]. The date must be a real calendar date."
        ) from exc
    if parsed.strftime("%Y%m%d") != date_mon:
        raise ConfigError(
            "date_Mon 不是一个真实日期 / "
            "Invalid date_Mon in [settings]. The date must be a real calendar date."
        )
    if not overview:
        raise ConfigError(
            "缺少 [overview_diseases] 中的疾病列表 / Missing diseases in [overview_diseases]."
        )
    if not focus:
        raise ConfigError(
            "缺少 [focus_diseases] 中的疾病列表 / Missing diseases in [focus_diseases]."
        )

    return ReportConfig(date_mon, overview, focus)


def write_report_config(path: Path, config: ReportConfig) -> None:
    content = CONFIG_TEMPLATE.format(
        date_mon=config.date_mon,
        overview_diseases="\n".join(config.overview_diseases),
        focus_diseases="\n".join(config.focus_diseases),
    )
    path.write_text(content, encoding="utf-8-sig", newline="\n")


def duplicate_values(values: list[str]) -> list[str]:
    seen: set[str] = set()
    duplicates: list[str] = []
    for value in values:
        if value in seen and value not in duplicates:
            duplicates.append(value)
        seen.add(value)
    return duplicates


def check_duplicates(values: list[str], section: str) -> None:
    duplicates = duplicate_values(values)
    if duplicates:
        raise ConfigError(
            f"[{section}] 中有重复疾病名称 / Duplicate disease names in [{section}]: "
            f"{join_zh(duplicates)}."
        )


def unique_headers(headers: list[Any]) -> list[str]:
    values = ["" if header is None else str(header) for header in headers]
    counts: dict[str, int] = {}
    for value in values:
        counts[value] = counts.get(value, 0) + 1

    result: list[str] = []
    for index, value in enumerate(values, start=1):
        if counts[value] > 1:
            result.append(f"{value}...{index}")
        else:
            result.append(value)
    return result


def read_excel_table(path: Path) -> tuple[list[dict[str, Any]], list[str]]:
    if not path.exists():
        raise FileNotFoundError(f"找不到输入文件 / Missing input file: {path}")
    workbook = load_workbook(str(path))
    worksheet = workbook.get_sheet_by_name("Report") if "Report" in workbook.sheet_names else workbook.get_sheet_by_index(0)
    rows = worksheet.to_python()

    if len(rows) < 2:
        raise ValueError(f"Excel 文件格式不正确 / Invalid Excel file format: {path}")

    headers = unique_headers(list(rows[1]))
    records: list[dict[str, Any]] = []
    for values in rows[2:]:
        record = {header: value for header, value in zip(headers, values)}
        if any(value is not None for value in values):
            records.append(record)
    return records, headers


def read_dictionary(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        raise FileNotFoundError(f"找不到疾病字典 / Missing dictionary file: {path}")
    with path.open("r", encoding="gbk", newline="") as handle:
        return [dict(row) for row in csv.DictReader(handle)]


def split_raw_extra(rows: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    for index, row in enumerate(rows):
        if strip_spaces(row.get("疾病病种")) == "其他传染病":
            return rows[:index], rows[index:]
    return rows, []


def build_df_all(raw_rows: list[dict[str, Any]], dictionary_rows: list[dict[str, str]]) -> list[dict[str, Any]]:
    dict_by_raw = {row["疾病病种"]: row for row in dictionary_rows}
    result: list[dict[str, Any]] = []
    for row in raw_rows:
        clean_name = strip_spaces(row.get("疾病病种"))
        dict_row = dict_by_raw.get(clean_name)
        if dict_row is None:
            continue
        if dict_row.get("tier") not in {"甲类", "乙类", "丙类"}:
            continue
        merged = dict(row)
        merged["疾病病种"] = clean_name
        merged["diseases"] = dict_row.get("diseases", "")
        merged["tier"] = dict_row.get("tier", "")
        merged["type"] = dict_row.get("type", "")
        for column in [
            "本期发病数",
            "上期发病数",
            "去年同期发病数",
            "本年至本期累计发病数",
            "去年至本期累计发病数",
            "本期死亡数",
            "上期死亡数",
            "去年同期死亡数",
            "本年至本期累计死亡数",
            "去年至本期累计死亡数",
        ]:
            merged[column] = as_number(merged.get(column))
        result.append(merged)
    return result


def clean_age_table(
    age_rows: list[dict[str, Any]], age_headers: list[str], dictionary_rows: list[dict[str, str]]
) -> tuple[list[dict[str, Any]], list[str]]:
    cleaned_headers = [header.replace(" ", "") for header in age_headers]
    n_tier_names = [row["疾病病种"] for row in dictionary_rows if row.get("tier") == "N"]
    keep_indices = [
        index
        for index, header in enumerate(cleaned_headers)
        if not any(name and name in header for name in n_tier_names)
    ]
    result_headers = [cleaned_headers[index] for index in keep_indices]
    result_rows: list[dict[str, Any]] = []
    for row in age_rows:
        values = list(row.values())
        result_rows.append(
            {
                result_headers[target_index]: values[source_index] if source_index < len(values) else None
                for target_index, source_index in enumerate(keep_indices)
            }
        )
    return result_rows, result_headers


def validate_config(
    config: ReportConfig,
    dictionary_rows: list[dict[str, str]],
    df_all: list[dict[str, Any]],
    age_headers: list[str],
) -> None:
    encoding_note = (
        " 请检查疾病名称拼写，并确认 config/report_config.txt 保存为 UTF-8 编码 / "
        "Please check disease spelling and make sure config/report_config.txt is saved as UTF-8."
    )
    check_duplicates(config.overview_diseases, "overview_diseases")
    check_duplicates(config.focus_diseases, "focus_diseases")

    dict_diseases = {row["diseases"] for row in dictionary_rows}
    all_diseases = {row["diseases"] for row in df_all}

    missing_overview_dict = [value for value in config.overview_diseases if value not in dict_diseases]
    if missing_overview_dict:
        raise ConfigError(
            "[overview_diseases] 中存在未知疾病名称 / Unknown disease names in [overview_diseases]: "
            f"{join_zh(missing_overview_dict)}.{encoding_note}"
        )

    missing_overview_data = [value for value in config.overview_diseases if value not in all_diseases]
    if missing_overview_data:
        raise ConfigError(
            "[overview_diseases] 中的疾病未出现在 B.xlsx 主分析数据中 / "
            "Disease names in [overview_diseases] are not available in the main B.xlsx analysis data: "
            f"{join_zh(missing_overview_data)}.{encoding_note}"
        )

    missing_focus_data = [value for value in config.focus_diseases if value not in all_diseases]
    if missing_focus_data:
        raise ConfigError(
            "[focus_diseases] 中的疾病未出现在 B.xlsx 主分析数据中 / "
            "Disease names in [focus_diseases] are not available in the main B.xlsx analysis data: "
            f"{join_zh(missing_focus_data)}.{encoding_note}"
        )

    focus_dict = [row for row in dictionary_rows if row["diseases"] in config.focus_diseases]
    focus_dict_diseases = {row["diseases"] for row in focus_dict}
    missing_focus_dict = [value for value in config.focus_diseases if value not in focus_dict_diseases]
    if missing_focus_dict:
        raise ConfigError(
            "[focus_diseases] 中存在未知疾病名称 / Unknown disease names in [focus_diseases]: "
            f"{join_zh(missing_focus_dict)}.{encoding_note}"
        )

    missing_age = [
        row["diseases"]
        for row in focus_dict
        if not any(row["疾病病种"] in header for header in age_headers)
    ]
    if missing_age:
        raise ConfigError(
            "[focus_diseases] 中的疾病未出现在 A.xlsx 年龄数据中 / "
            "Disease names in [focus_diseases] are not available in A.xlsx age data: "
            f"{join_zh(missing_age)}.{encoding_note}"
        )


def rate_calc(later: Any, former: Any) -> str:
    later_num = as_number(later)
    former_num = as_number(former)
    diff = later_num - former_num
    if former_num == 0:
        if later_num == 0:
            return "一致"
        return f"增加{fmt_int(diff)}例"
    rate = diff / former_num * 100
    if rate > 0:
        return f"上升{fmt_pct(abs(rate))}%"
    if rate == 0:
        return "一致"
    return f"下降{fmt_pct(abs(rate))}%"


def dist_calc(later: Any, former: Any) -> str:
    later_num = as_number(later)
    former_num = as_number(former)
    diff = later_num - former_num
    if former_num == 0 or diff == 0:
        return ""
    if diff > 0:
        return f"（增加{fmt_int(abs(diff))}例）"
    return f"（减少{fmt_int(abs(diff))}例）"


def ratedisease(disease: str, df_all: list[dict[str, Any]]) -> dict[str, str]:
    row = next(item for item in df_all if item["diseases"] == disease)
    count_l = row["本期发病数"]
    count_f = row["上期发病数"]
    r = rate_calc(count_l, count_f)
    dist = dist_calc(count_l, count_f)
    sen_1 = f"{disease}较上周{r}{dist}"
    sen_2 = (
        f"本周报告{fmt_int(count_l)}例，较上周（{fmt_int(count_f)}例）{r}{dist}，"
        f"较去年同期（{fmt_int(row['去年同期发病数'])}例）"
        f"{rate_calc(count_l, row['去年同期发病数'])}"
        f"{dist_calc(count_l, row['去年同期发病数'])}，"
    )
    sen_3 = (
        f"本年度累计报告{fmt_int(row['本年至本期累计发病数'])}例，"
        f"较去年同期（{fmt_int(row['去年至本期累计发病数'])}例）"
        f"{rate_calc(row['本年至本期累计发病数'], row['去年至本期累计发病数'])}"
        f"{dist_calc(row['本年至本期累计发病数'], row['去年至本期累计发病数'])}。"
    )
    return {"tier": row["tier"], "sen_1": sen_1, "sen_2": sen_2, "sen_3": sen_3}


def overview_tier(tier: str, df_all: list[dict[str, Any]]) -> dict[str, Any]:
    rows = [row for row in df_all if row["tier"] == tier]
    count_l = sum(row["本期发病数"] for row in rows)
    count_f = sum(row["上期发病数"] for row in rows)
    class_count = sum(1 for row in rows if row["本期发病数"] != 0)
    if count_l == 0:
        sentence = f"{tier}传染病无病例报告。"
    else:
        sentence = (
            f"{tier}传染病报告{class_count}种{fmt_int(count_l)}例，"
            f"发病数较上周（{fmt_int(count_f)}例）{rate_calc(count_l, count_f)}；"
        )
    return {"tier": tier, "count_l": count_l, "sen": sentence}


def merge_tier(tier: str, tier_summaries: dict[str, dict[str, Any]], disease_summaries: list[dict[str, str]]) -> str:
    tier_summary = tier_summaries[tier]
    tier_diseases = [item for item in disease_summaries if item["tier"] == tier]
    if tier_summary["count_l"] == 0:
        return tier_summary["sen"]
    return f"{tier_summary['sen']}其中{join_zh([item['sen_1'] for item in tier_diseases], '，')}。"


def parse_age(value: Any) -> float | None:
    if value is None:
        return None
    text = str(value).replace("-", "").replace("及以上", "").replace("不详", "-1").strip()
    if text == "":
        return None
    return float(text)


def age_calc(label: str, rows: list[dict[str, float]]) -> str:
    n_all = sum(row["n"] for row in rows)
    if label == "5岁及以下":
        filtered = [row for row in rows if 0 <= row["age"] <= 5]
    elif label == "6-19岁":
        filtered = [row for row in rows if 6 <= row["age"] <= 19]
    elif label == "20-59岁":
        filtered = [row for row in rows if 20 <= row["age"] <= 59]
    else:
        filtered = [row for row in rows if row["age"] >= 60]
    n_d = sum(row["n"] for row in filtered)
    if n_d != 0:
        return f"{label}报告{fmt_int(n_d)}例（占{fmt_pct(n_d / n_all * 100)}%）"
    return f"{label}报告{fmt_int(n_d)}例"


def disease_focus(
    disease: str,
    age_rows: list[dict[str, Any]],
    age_headers: list[str],
    df_all: list[dict[str, Any]],
) -> dict[str, Any]:
    main_row = next(row for row in df_all if row["diseases"] == disease)
    raw_disease = main_row["疾病病种"]
    selected_headers = [
        header for header in age_headers if "疾病病种" in header or raw_disease in header
    ]
    if len(selected_headers) < 4:
        raise ConfigError(
            f"[focus_diseases] 中的疾病未出现在 A.xlsx 年龄数据中 / "
            f"Disease names in [focus_diseases] are not available in A.xlsx age data: {disease}."
        )
    age_header = selected_headers[0]
    count_header = selected_headers[3]

    parsed_rows: list[dict[str, float]] = []
    for row in age_rows[3:]:
        age = parse_age(row.get(age_header))
        if age is None:
            continue
        parsed_rows.append({"age": age, "n": as_number(row.get(count_header))})
    n = sum(row["n"] for row in parsed_rows)
    sen_age = "，".join(
        age_calc(label, parsed_rows)
        for label in ["5岁及以下", "6-19岁", "20-59岁", "60岁及以上"]
    )
    disease_sens = ratedisease(disease, df_all)
    return {
        "v": disease,
        "tier": main_row["tier"],
        "n": n,
        "sen_1": f"{disease_sens['sen_2']}周病例数趋势见图",
        "sen_2": f"本周报告病例中，{sen_age}。",
        "sen_3": disease_sens["sen_3"],
    }


def other_diseases(extra_rows: list[dict[str, Any]]) -> list[str]:
    rows: list[dict[str, Any]] = []
    for index, row in enumerate(extra_rows):
        current = as_number(row.get("本期发病数"))
        if current == 0:
            continue
        rows.append(
            {
                **row,
                "row_id": index,
                "疾病病种": strip_spaces(row.get("疾病病种")),
                "本期发病数": current,
                "上期发病数": as_number(row.get("上期发病数")),
                "本年至本期累计发病数": as_number(row.get("本年至本期累计发病数")),
                "去年至本期累计发病数": as_number(row.get("去年至本期累计发病数")),
            }
        )
    rows.sort(key=lambda row: (-row["本期发病数"], row["row_id"]))
    return [
        (
            f"{row['疾病病种']}{fmt_int(row['本期发病数'])}例，"
            f"较上周（{fmt_int(row['上期发病数'])}例）"
            f"{rate_calc(row['本期发病数'], row['上期发病数'])}，"
            f"本年度累计报告{fmt_int(row['本年至本期累计发病数'])}例，"
            f"较去年同期（{fmt_int(row['去年至本期累计发病数'])}例）"
            f"{rate_calc(row['本年至本期累计发病数'], row['去年至本期累计发病数'])}。"
        )
        for row in rows
    ]


def generate_report(root: Path | None = None) -> ReportResult:
    root = root or project_root()
    config = read_report_config(root / "config" / "report_config.txt")
    raw_age_rows, raw_age_headers = read_excel_table(root / "input_database" / "A.xlsx")
    raw_rows, _raw_headers = read_excel_table(root / "input_database" / "B.xlsx")
    dictionary_rows = read_dictionary(root / "config" / "dictionary.csv")

    main_raw_rows, raw_extra_rows = split_raw_extra(raw_rows)
    df_all = build_df_all(main_raw_rows, dictionary_rows)
    age_rows, age_headers = clean_age_table(raw_age_rows, raw_age_headers, dictionary_rows)
    validate_config(config, dictionary_rows, df_all, age_headers)

    date_mon = datetime.strptime(config.date_mon, "%Y%m%d")
    week = date_mon.isocalendar().week
    year = date_mon.year

    count_all = sum(row["本期发病数"] for row in df_all)
    class_all = sum(1 for row in df_all if row["本期发病数"] != 0)
    for row in df_all:
        row["rate"] = row["本期发病数"] / count_all

    death_rows = [
        row
        for row in sorted(df_all, key=lambda item: -item["本期死亡数"])
        if row["本期死亡数"] != 0
    ]
    count_death = sum(row["本期死亡数"] for row in death_rows)
    death_text = ""
    if death_rows:
        death_text = "（" + "，".join(
            f"{row['diseases']}{fmt_int(row['本期死亡数'])}例" for row in death_rows
        ) + "）"

    fifth_sorted = sorted(df_all, key=lambda item: -item["本期发病数"])
    fifth_threshold = fifth_sorted[4]["本期发病数"]
    count_fifth = [row for row in fifth_sorted if row["本期发病数"] >= fifth_threshold]
    rate_fifth = sum(row["rate"] for row in count_fifth) * 100
    first_raw = main_raw_rows[0]
    p1_1 = (
        f"{year}年第{week}周全市共报告法定传染病{class_all}种{fmt_int(count_all)}例，"
        f"死亡{fmt_int(count_death)}例{death_text}，详见附表。"
        f"报告病例数居前{len(count_fifth)}位的病种依次为："
        f"{join_zh([f'{row["diseases"]}（{fmt_int(row["本期发病数"])}例）' for row in count_fifth])}，"
        f"共占法定传染病报告发病数的{fmt_pct(rate_fifth)}%。"
        f"与上周（{fmt_int(first_raw['上期发病数'])}例）相比传染病报告发病数"
        f"{rate_calc(first_raw['本期发病数'], first_raw['上期发病数'])}。"
    )

    tier_summaries = {tier: overview_tier(tier, df_all) for tier in ["甲类", "乙类", "丙类"]}
    overview_diseases = [ratedisease(disease, df_all) for disease in config.overview_diseases]
    p1_2 = "".join(
        merge_tier(tier, tier_summaries, overview_diseases)
        for tier in ["甲类", "乙类", "丙类"]
    )

    tier_order = {"甲类": 0, "乙类": 1, "丙类": 2}
    focus_rows = [
        disease_focus(disease, age_rows, age_headers, df_all)
        for disease in config.focus_diseases
    ]
    focus_rows.sort(key=lambda row: (tier_order[row["tier"]], -row["n"]))
    sens_2: list[tuple[str, list[str]]] = []
    for index, row in enumerate(focus_rows, start=1):
        if index > len(SECTION_LABELS):
            raise ConfigError(f"Missing section label for order {index}.")
        title = f"{SECTION_LABELS[index - 1]}{row['v']}"
        sentence = f"{row['sen_1']}{index}。{row['sen_2']}{row['sen_3']}"
        sens_2.append((title, [sentence]))

    other_sentences = other_diseases(raw_extra_rows)
    if other_sentences:
        other_order = len(sens_2) + 1
        if other_order > len(SECTION_LABELS):
            raise ConfigError(f"Missing section label for order {other_order}.")
        sens_2.append((f"{SECTION_LABELS[other_order - 1]}其他传染病", other_sentences))

    report_change = (as_number(first_raw["本期发病数"]) - as_number(first_raw["上期发病数"])) / as_number(first_raw["上期发病数"]) * 100
    if report_change >= 10:
        trend = "明显上升"
    elif report_change >= 5:
        trend = "略有上升"
    elif report_change > -5:
        trend = "基本持平"
    elif report_change >= -10:
        trend = "略有下降"
    else:
        trend = "明显下降"

    respiratory_count = sum(row["本期发病数"] for row in df_all if row["type"] == "呼吸道传染病")
    respiratory_rate = sum(row["rate"] for row in df_all if row["type"] == "呼吸道传染病") * 100
    intestinal_count = sum(row["本期发病数"] for row in df_all if row["type"] == "肠道传染病")
    intestinal_rate = sum(row["rate"] for row in df_all if row["type"] == "肠道传染病") * 100
    p3_1 = (
        f"本周我市法定传染病报告发病数较上周{trend}（{fmt_pct(report_change)}%），"
        f"其中呼吸道传染病（新型冠状病毒感染、流行性感冒、肺结核、百日咳、猩红热等）"
        f"报告{fmt_int(respiratory_count)}例（占{respiratory_rate:.0f}%），"
        f"肠道传染病（痢疾、手足口病、其他感染性腹泻病等）"
        f"报告{fmt_int(intestinal_count)}例（占{intestinal_rate:.0f}%）。"
    )

    table_rows = [["病名", "发病数", "死亡数"]]
    report_table_rows = [
        row
        for row in sorted(df_all, key=lambda item: -item["本期发病数"])
        if row["本期发病数"] != 0
    ]
    table_rows.append(["甲乙丙类总计", int(count_all), int(count_death)])
    table_rows.extend(
        [row["diseases"], int(row["本期发病数"]), int(row["本期死亡数"])]
        for row in report_table_rows
    )

    paragraphs: list[str] = [
        "一、疫情概况",
        p1_1,
        p1_2,
        "\n",
        "二、重点疫情",
    ]
    for title, sentences in sens_2:
        paragraphs.append(title)
        paragraphs.extend(sentences)
    paragraphs.extend(["\n", "三、重点疫情健康提示", p3_1])

    outputs = root / "outputs"
    outputs.mkdir(exist_ok=True)
    docx_path = outputs / "output_beta.docx"
    csv_path = outputs / "table_beta.csv"
    write_docx(docx_path, paragraphs)
    write_table_csv(csv_path, table_rows)
    return ReportResult(docx_path, csv_path, paragraphs, table_rows)


def write_docx(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    _set_document_fonts(document)
    for paragraph in paragraphs:
        doc_paragraph = document.add_paragraph()
        if paragraph == "\n":
            run = doc_paragraph.add_run()
            text = OxmlElement("w:t")
            text.set(qn("xml:space"), "preserve")
            text.text = "\n"
            run._r.append(text)
        else:
            run = doc_paragraph.add_run(paragraph)
        _set_run_fonts(run)
    document.save(path)


def _set_document_fonts(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    r_pr = style.element.get_or_add_rPr()
    _set_rfonts(r_pr)


def _set_run_fonts(run: Any) -> None:
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    _set_rfonts(r_pr)


def _set_rfonts(r_pr: Any) -> None:
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")


def write_table_csv(path: Path, table_rows: list[list[Any]]) -> None:
    lines: list[str] = []
    for index, row in enumerate(table_rows):
        if index == 0:
            lines.append(",".join(f'"{str(value).replace(chr(34), chr(34) + chr(34))}"' for value in row))
            continue
        first = str(row[0]).replace('"', '""')
        lines.append(f'"{first}",{row[1]},{row[2]}')
    path.write_text("\n".join(lines) + "\n", encoding="gbk", newline="")


def load_ui_options(root: Path | None = None) -> tuple[ReportConfig, list[str]]:
    root = root or project_root()
    config = read_report_config(root / "config" / "report_config.txt")
    dictionary_rows = read_dictionary(root / "config" / "dictionary.csv")
    disease_options = [
        row["diseases"]
        for row in dictionary_rows
        if row.get("tier") in {"甲类", "乙类", "丙类"}
    ]
    for disease in config.overview_diseases + config.focus_diseases:
        if disease not in disease_options:
            disease_options.append(disease)
    return config, disease_options


if __name__ == "__main__":
    result = generate_report(project_root())
    print(f"Generated: {result.docx_path}")
    print(f"Generated: {result.csv_path}")
